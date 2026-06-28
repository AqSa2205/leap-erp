import io
import os
import re
import copy
import zipfile
from lxml import etree
from django.http import HttpResponse
from django.conf import settings
from django.contrib.staticfiles import finders
from django.utils.text import slugify

WNS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
WPS_NS = 'http://schemas.microsoft.com/office/word/2010/wordprocessingShape'
VML_NS = 'urn:schemas-microsoft-com:vml'
DRAWING_NS = 'http://schemas.openxmlformats.org/drawingml/2006/main'
WP_NS = 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing'
PIC_NS = 'http://schemas.openxmlformats.org/drawingml/2006/picture'
REL_NS = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'
XML_SPACE = '{http://www.w3.org/XML/1998/namespace}space'

# Detect rich-text HTML (from TinyMCE / Word paste) regardless of tag
# attributes. Matches opening/closing tags of known HTML elements with a word
# boundary, so a stray "<" in plain text won't false-positive.
_HTML_TAG_RE = re.compile(
    r'</?(?:p|div|span|strong|b|em|i|u|s|br|hr|table|tr|td|th|thead|tbody|'
    r'ul|ol|li|img|figure|figcaption|h[1-6]|a|sub|sup|blockquote|pre|code)\b',
    re.IGNORECASE)


def _find_template():
    """Locate the DOCX template file."""
    result = finders.find('docx_templates/technical_proposal_template.docx')
    if result:
        return result
    path = os.path.join(settings.BASE_DIR, 'static', 'docx_templates',
                        'technical_proposal_template.docx')
    if os.path.exists(path):
        return path
    return None


# ── Strip yellow highlights ──────────────────────────────────

def _strip_highlights(root):
    """Remove all w:highlight elements from the XML tree."""
    for highlight in root.findall(f'.//{{{WNS}}}highlight'):
        highlight.getparent().remove(highlight)


# ── Text-box run-merging replacement ─────────────────────────

def _replace_in_textbox_content(txbx_content, replacements, exact_replacements):
    """
    Join all w:t text in a txbxContent element, apply replacements,
    then put the result in the first w:t and clear the rest.

    replacements: dict of substring replacements (applied with str.replace)
    exact_replacements: dict where key must match the ENTIRE joined text
                        (for short values like initials that would otherwise
                         match inside longer words)
    """
    t_elements = txbx_content.findall(f'.//{{{WNS}}}t')
    if not t_elements:
        return

    joined = ''.join((t.text or '') for t in t_elements)
    original = joined

    # Try exact match first (for initials)
    stripped = joined.strip()
    if stripped in exact_replacements:
        joined = exact_replacements[stripped]
    else:
        # Substring replacements
        for old, new in replacements.items():
            joined = joined.replace(old, new)

    if joined == original:
        return  # nothing changed

    # Put all text in first w:t, clear the rest
    t_elements[0].text = joined
    t_elements[0].set(XML_SPACE, 'preserve')
    for t in t_elements[1:]:
        t.text = ''


def _replace_textboxes_in_part(root, replacements, exact_replacements):
    """
    Find all text boxes (WPS and VML) in an XML part and apply
    run-merging replacements to each.
    """
    # WPS text boxes: wps:txbx > w:txbxContent
    for txbx in root.iter(f'{{{WPS_NS}}}txbx'):
        txbx_content = txbx.find(f'{{{WNS}}}txbxContent')
        if txbx_content is not None:
            _replace_in_textbox_content(txbx_content, replacements, exact_replacements)

    # VML fallback text boxes: v:textbox > w:txbxContent
    for textbox in root.iter(f'{{{VML_NS}}}textbox'):
        txbx_content = textbox.find(f'{{{WNS}}}txbxContent')
        if txbx_content is not None:
            _replace_in_textbox_content(txbx_content, replacements, exact_replacements)


# ── Paragraph run-merging replacement ────────────────────────

def _replace_in_paragraph_runs(para, replacements):
    """
    Join all w:t in a paragraph's runs, apply replacements,
    put result in first w:t, clear the rest.
    """
    runs = para.findall(f'{{{WNS}}}r')
    if not runs:
        return

    t_elements = []
    for r in runs:
        t = r.find(f'{{{WNS}}}t')
        if t is not None:
            t_elements.append(t)

    if not t_elements:
        return

    joined = ''.join((t.text or '') for t in t_elements)
    original = joined

    for old, new in replacements.items():
        joined = joined.replace(old, new)

    if joined == original:
        return

    t_elements[0].text = joined
    t_elements[0].set(XML_SPACE, 'preserve')
    for t in t_elements[1:]:
        t.text = ''


# ── Cover page replacement ───────────────────────────────────

def _replace_cover_page(body, proposal):
    """
    Replace cover page text in the first ~25 paragraphs.
    """
    paragraphs = body.findall(f'{{{WNS}}}p')

    cover_replacements = {}

    # P[1]: reference (template has lowercase-L typo "LNUK-IRLl02125070")
    cover_replacements['LNUK-IRLl02125070'] = proposal.proposal_reference
    cover_replacements['LNUK-IRL02125070'] = proposal.proposal_reference

    # P[2]: project description
    cover_replacements['PROVISIONING OF CCTV, IIS, ACS AND FTTH SYSTEM SOLUTION'] = (
        proposal.project_description.upper() if proposal.project_description else ''
    )

    # P[4]: site/location -> client name
    cover_replacements['CANAL SIDE MANCHESTER'] = proposal.client_name.upper()

    # P[6]: document type
    cover_replacements['PRELIMINARY TECHNICAL PROPOSAL'] = proposal.document_type.upper()

    # Client name variants (on cover page "Prepared for:" area)
    cover_replacements['MERIDIAM Construction'] = proposal.client_name

    # Apply longest matches first to avoid partial double-replacement
    sorted_replacements = dict(
        sorted(cover_replacements.items(), key=lambda x: len(x[0]), reverse=True)
    )

    for i in range(min(25, len(paragraphs))):
        _replace_in_paragraph_runs(paragraphs[i], sorted_replacements)


# ── Body section content replacement ─────────────────────────

def _get_paragraph_text(para):
    return ''.join(
        (t.text or '') for t in para.findall(f'.//{{{WNS}}}t')
    ).strip()


def _get_paragraph_style(para):
    pPr = para.find(f'{{{WNS}}}pPr')
    if pPr is not None:
        pStyle = pPr.find(f'{{{WNS}}}pStyle')
        if pStyle is not None:
            return pStyle.get(f'{{{WNS}}}val', '')
    return ''


def _extract_content_formatting(source_para):
    """
    Extract paragraph properties (pPr) and run properties (rPr)
    from a content paragraph to use as a template for new paragraphs.
    Returns (pPr_element_or_None, rPr_element_or_None).
    """
    pPr = None
    source_pPr = source_para.find(f'{{{WNS}}}pPr')
    if source_pPr is not None:
        pPr = copy.deepcopy(source_pPr)
        # Remove heading style so new paragraphs aren't styled as headings
        pStyle = pPr.find(f'{{{WNS}}}pStyle')
        if pStyle is not None:
            pPr.remove(pStyle)

    rPr = None
    first_run = source_para.find(f'{{{WNS}}}r')
    if first_run is not None:
        rPr_el = first_run.find(f'{{{WNS}}}rPr')
        if rPr_el is not None:
            rPr = copy.deepcopy(rPr_el)

    return pPr, rPr


def _make_default_pPr():
    """
    Create default paragraph properties matching the template's content style:
    left indent 709 twips, line spacing 276 (1.15), justify both,
    Trebuchet MS 10pt.
    """
    pPr = etree.Element(f'{{{WNS}}}pPr')
    spacing = etree.SubElement(pPr, f'{{{WNS}}}spacing')
    spacing.set(f'{{{WNS}}}line', '276')
    spacing.set(f'{{{WNS}}}lineRule', 'auto')
    ind = etree.SubElement(pPr, f'{{{WNS}}}ind')
    ind.set(f'{{{WNS}}}left', '709')
    jc = etree.SubElement(pPr, f'{{{WNS}}}jc')
    jc.set(f'{{{WNS}}}val', 'both')
    return pPr


def _make_default_rPr():
    """Default run properties: Trebuchet MS 10pt."""
    rPr = etree.Element(f'{{{WNS}}}rPr')
    rFonts = etree.SubElement(rPr, f'{{{WNS}}}rFonts')
    rFonts.set(f'{{{WNS}}}ascii', 'Trebuchet MS')
    rFonts.set(f'{{{WNS}}}hAnsi', 'Trebuchet MS')
    sz = etree.SubElement(rPr, f'{{{WNS}}}sz')
    sz.set(f'{{{WNS}}}val', '20')
    szCs = etree.SubElement(rPr, f'{{{WNS}}}szCs')
    szCs.set(f'{{{WNS}}}val', '20')
    return rPr


def _make_content_paragraph(text, pPr_template=None, rPr_template=None):
    """Create a new w:p element with the given text, paragraph and run formatting."""
    p = etree.Element(f'{{{WNS}}}p')

    # Paragraph properties (indentation, spacing, justification)
    if pPr_template is not None:
        p.insert(0, copy.deepcopy(pPr_template))
    else:
        p.insert(0, _make_default_pPr())

    r = etree.SubElement(p, f'{{{WNS}}}r')

    # Run properties (font, size)
    if rPr_template is not None:
        r.insert(0, copy.deepcopy(rPr_template))
    else:
        r.insert(0, _make_default_rPr())

    t = etree.SubElement(r, f'{{{WNS}}}t')
    t.text = text
    t.set(XML_SPACE, 'preserve')
    return p


def _is_heading1(elem):
    """Check if a body child element is a Heading1 paragraph."""
    if elem.tag != f'{{{WNS}}}p':
        return False
    return _get_paragraph_style(elem) == 'Heading1'


def _set_heading_text(para, text):
    """Set a heading paragraph's visible text, preserving run formatting and
    stripping bookmark elements (so cloned headings don't duplicate IDs)."""
    # Drop TOC bookmarks — duplicating their w:id across cloned headings
    # would corrupt the document.
    for tag in ('bookmarkStart', 'bookmarkEnd'):
        for bm in para.findall(f'{{{WNS}}}{tag}'):
            para.remove(bm)

    t_elements = para.findall(f'.//{{{WNS}}}t')
    if t_elements:
        t_elements[0].text = text
        t_elements[0].set(XML_SPACE, 'preserve')
        for t in t_elements[1:]:
            t.text = ''
    else:
        r = etree.SubElement(para, f'{{{WNS}}}r')
        r.insert(0, _make_default_rPr())
        t = etree.SubElement(r, f'{{{WNS}}}t')
        t.text = text
        t.set(XML_SPACE, 'preserve')


def _replace_body_sections(xml_bytes, proposal, image_registry):
    """Rebuild the proposal body from the user's ProposalSection rows.

    The template carries a fixed set of Heading1 sections. We keep everything
    before the first Heading1 (cover page + table of contents) and the trailing
    sectPr, wipe the old section region entirely, then emit one Heading1 +
    content block per ProposalSection (in order), reusing the template's
    Heading1 styling. Engineering documents, if any, are appended as a final
    section using the template's documents table.

    image_registry: mutable list populated with image metadata as <img>
    tags are encountered. Used to later add image files and relationships
    to the ZIP.
    """
    root = etree.fromstring(xml_bytes)
    body = root.find(f'.//{{{WNS}}}body')
    if body is None:
        return xml_bytes

    children = list(body)

    # Find Heading1 positions — these delimit the template's section region.
    heading_positions = [i for i, c in enumerate(children) if _is_heading1(c)]
    if not heading_positions:
        return xml_bytes
    first_heading_idx = heading_positions[0]

    # Heading template: reuse the first Heading1 paragraph's formatting.
    heading_template = copy.deepcopy(children[first_heading_idx])

    # Content formatting template: first non-empty plain content paragraph
    # following any heading (Heading2/list paragraphs are skipped).
    pPr_template = None
    rPr_template = None
    for hpos in heading_positions:
        for ci in range(hpos + 1, min(hpos + 4, len(children))):
            cand = children[ci]
            if (cand.tag == f'{{{WNS}}}p' and not _get_paragraph_style(cand)
                    and _get_paragraph_text(cand)):
                pPr_template, rPr_template = _extract_content_formatting(cand)
                break
        if pPr_template is not None:
            break

    # Capture the engineering documents table (before we wipe the region) so we
    # can re-emit it with the user's data after the sections.
    eng_table_template = None
    for tbl in body.findall(f'.//{{{WNS}}}tbl'):
        header = ''.join(
            (t.text or '') for t in tbl.findall(f'.//{{{WNS}}}t')[:8]
        ).lower()
        if 'document' in header:
            eng_table_template = copy.deepcopy(tbl)
            break

    # Trailing sectPr boundary — preserve page/section setup at the end.
    end_idx = len(children)
    while end_idx > first_heading_idx and children[end_idx - 1].tag == f'{{{WNS}}}sectPr':
        end_idx -= 1

    # Anchor = last cover-page element we keep; new content is inserted after it.
    anchor = children[first_heading_idx - 1]

    # Wipe the old section region (headings + content + leftover template junk).
    for j in range(end_idx - 1, first_heading_idx - 1, -1):
        body.remove(children[j])

    # Build the new section blocks from the user's rows.
    new_elements = []
    for section in proposal.sections.all():
        heading_para = copy.deepcopy(heading_template)
        _set_heading_text(heading_para, section.heading)
        new_elements.append(heading_para)

        content = section.content or ''
        if not content.strip():
            continue
        if _HTML_TAG_RE.search(content):
            new_elements.extend(
                _html_to_elements(content, pPr_template, rPr_template, image_registry))
        else:
            for line in content.split('\n'):
                new_elements.append(
                    _make_content_paragraph(line, pPr_template, rPr_template))

    # Engineering documents as a final section.
    eng_docs = list(proposal.engineering_documents.all())
    if eng_docs and eng_table_template is not None:
        heading_para = copy.deepcopy(heading_template)
        _set_heading_text(heading_para, 'Engineering Documents')
        new_elements.append(heading_para)
        new_elements.append(copy.deepcopy(eng_table_template))

    # Insert everything after the cover region, before the trailing sectPr.
    current = anchor
    for el in new_elements:
        current.addnext(el)
        current = el

    # Rebuild the engineering documents table rows from user data (operates in
    # place on the table we just inserted).
    _fill_engineering_table(body, proposal)

    return etree.tostring(root, xml_declaration=True, encoding='UTF-8', standalone=True)


def _html_to_elements(html_content, pPr_template, rPr_template, image_registry):
    """Convert HTML from TinyMCE to a list of DOCX elements (paragraphs, tables,
    images). The caller is responsible for inserting them into the body.

    image_registry is a mutable list the parser appends (src, alt, width, height)
    tuples to as it encounters <img> tags. Each entry's index becomes the
    image's rId suffix (rIdImg0, rIdImg1, ...) that we add to the zip later.
    """
    import re
    from html import unescape
    from html.parser import HTMLParser

    class DocxHTMLParser(HTMLParser):
        """Parse HTML and build a list of DOCX elements (w:p and w:tbl)."""

        def __init__(self):
            super().__init__()
            self.elements = []        # final list of etree elements
            self._text = ''           # current accumulated text
            self._in_list = False
            self._list_type = None    # 'ul' or 'ol'
            self._list_counter = 0
            self._in_table = False
            self._table_rows = []     # list of rows, each row = list of cell texts
            self._current_row = []
            self._current_cell = ''
            self._is_header_row = False
            self._bold = False
            self._in_heading = None  # None or 'h2'/'h3'/'h4'
            self._in_figure = False
            self._figure_img = None    # dict with src, alt, width, height
            self._figure_caption = ''
            self._in_figcaption = False

        def _flush_text(self, prefix=''):
            text = unescape(self._text).strip()
            if text:
                p = _make_content_paragraph(prefix + text, pPr_template, rPr_template)
                self.elements.append(p)
            self._text = ''

        def _flush_heading(self):
            """Create a DOCX heading paragraph from accumulated text."""
            text = unescape(self._text).strip()
            if not text:
                self._text = ''
                return
            level = {'h2': 'Heading2', 'h3': 'Heading3', 'h4': 'Heading4'}.get(self._in_heading, 'Heading2')
            p = etree.Element(f'{{{WNS}}}p')
            pPr = etree.SubElement(p, f'{{{WNS}}}pPr')
            pStyle = etree.SubElement(pPr, f'{{{WNS}}}pStyle')
            pStyle.set(f'{{{WNS}}}val', level)
            r = etree.SubElement(p, f'{{{WNS}}}r')
            rPr = _make_default_rPr()
            bold = etree.SubElement(rPr, f'{{{WNS}}}b')
            r.insert(0, rPr)
            t = etree.SubElement(r, f'{{{WNS}}}t')
            t.text = text
            t.set(XML_SPACE, 'preserve')
            self.elements.append(p)
            self._text = ''

        def handle_starttag(self, tag, attrs):
            tag = tag.lower()
            attrs_dict = dict(attrs)
            if tag == 'p':
                self._text = ''
            elif tag in ('h2', 'h3', 'h4'):
                self._text = ''
                self._in_heading = tag
            elif tag == 'figure':
                self._in_figure = True
                self._figure_img = None
                self._figure_caption = ''
            elif tag == 'figcaption':
                self._in_figcaption = True
                self._figure_caption = ''
            elif tag == 'img':
                self._figure_img = {
                    'src': attrs_dict.get('src', ''),
                    'alt': attrs_dict.get('alt', ''),
                    'width': attrs_dict.get('width', ''),
                    'height': attrs_dict.get('height', ''),
                }
                # Bare <img> (no <figure>) — emit image paragraph immediately
                if not self._in_figure:
                    img_para = self._build_image_paragraph(self._figure_img, '')
                    if img_para is not None:
                        self.elements.append(img_para)
                    self._figure_img = None
            elif tag == 'ul':
                self._in_list = True
                self._list_type = 'ul'
                self._list_counter = 0
            elif tag == 'ol':
                self._in_list = True
                self._list_type = 'ol'
                self._list_counter = 0
            elif tag == 'li':
                self._text = ''
            elif tag == 'table':
                self._in_table = True
                self._table_rows = []
            elif tag == 'tr':
                self._current_row = []
                self._is_header_row = False
            elif tag == 'th':
                self._current_cell = ''
                self._is_header_row = True
            elif tag == 'td':
                self._current_cell = ''
            elif tag == 'br':
                self._text += '\n'
            elif tag in ('strong', 'b'):
                self._bold = True

        def handle_endtag(self, tag):
            tag = tag.lower()
            if tag in ('h2', 'h3', 'h4'):
                self._flush_heading()
                self._in_heading = None
            elif tag == 'figcaption':
                self._in_figcaption = False
            elif tag == 'figure':
                # Emit image paragraph + caption paragraph
                if self._figure_img:
                    img_para = self._build_image_paragraph(
                        self._figure_img, self._figure_caption
                    )
                    if img_para is not None:
                        self.elements.append(img_para)
                    if self._figure_caption:
                        cap_para = self._build_caption_paragraph(self._figure_caption)
                        self.elements.append(cap_para)
                self._in_figure = False
                self._figure_img = None
                self._figure_caption = ''
            elif tag == 'p':
                self._flush_text()
            elif tag in ('ul', 'ol'):
                self._in_list = False
            elif tag == 'li':
                self._list_counter += 1
                prefix = f'{self._list_counter}. ' if self._list_type == 'ol' else '\u2022 '
                self._flush_text(prefix)
            elif tag in ('td', 'th'):
                self._current_row.append(unescape(self._current_cell).strip())
                self._current_cell = ''
            elif tag == 'tr':
                self._table_rows.append((self._current_row, self._is_header_row))
            elif tag == 'table':
                self._in_table = False
                if self._table_rows:
                    self.elements.append(self._build_docx_table())
                self._table_rows = []
            elif tag in ('strong', 'b'):
                self._bold = False

        def handle_data(self, data):
            if self._in_figcaption:
                self._figure_caption += data
            elif self._in_table:
                self._current_cell += data
            else:
                self._text += data

        def _build_image_paragraph(self, img_info, alt_caption):
            """Register the image and build a w:p containing a w:drawing that references it.

            Uses raw XML string (not etree) to ensure proper namespace prefixes
            (pic:, a:, wp:, r:) that Word expects — lxml would otherwise emit
            generated prefixes like ns2: which some Word versions reject.
            """
            src = img_info.get('src', '')
            if not src:
                return None

            idx = len(image_registry)
            image_registry.append({
                'src': src,
                'alt': img_info.get('alt') or alt_caption or 'image',
                'width': img_info.get('width'),
                'height': img_info.get('height'),
            })
            # Placeholder rId — replaced with a real unique numeric rId when
            # we actually inject images into the zip
            rid_placeholder = f'__LEAP_IMG_{idx}__'

            # Force 600 x 300 pixel display size (9525 EMU per pixel at 96 DPI)
            cx = 600 * 9525   # 5,715,000 EMU ~ 15.9 cm
            cy = 300 * 9525   # 2,857,500 EMU ~ 7.9 cm

            pic_id = idx + 100  # start at 100 to avoid clashing with template pic ids
            alt = img_info.get('alt') or alt_caption or f'Picture {pic_id}'
            # Escape alt text for XML attribute
            alt_xml = (alt.replace('&', '&amp;').replace('"', '&quot;')
                          .replace('<', '&lt;').replace('>', '&gt;'))

            # Raw XML with proper prefixes. The parent document.xml already
            # declares all required namespaces (w, r, wp, a, pic), so the
            # prefixes resolve correctly without extra xmlns attributes.
            xml = (
                f'<w:p xmlns:w="{WNS}" xmlns:r="{REL_NS}" xmlns:wp="{WP_NS}" '
                f'xmlns:a="{DRAWING_NS}" xmlns:pic="{PIC_NS}">'
                f'<w:pPr><w:jc w:val="center"/></w:pPr>'
                f'<w:r><w:drawing>'
                f'<wp:inline distT="0" distB="0" distL="0" distR="0">'
                f'<wp:extent cx="{cx}" cy="{cy}"/>'
                f'<wp:effectExtent l="0" t="0" r="0" b="0"/>'
                f'<wp:docPr id="{pic_id}" name="Picture {pic_id}" descr="{alt_xml}"/>'
                f'<wp:cNvGraphicFramePr>'
                f'<a:graphicFrameLocks noChangeAspect="1"/>'
                f'</wp:cNvGraphicFramePr>'
                f'<a:graphic>'
                f'<a:graphicData uri="{PIC_NS}">'
                f'<pic:pic>'
                f'<pic:nvPicPr>'
                f'<pic:cNvPr id="{pic_id}" name="Picture {pic_id}"/>'
                f'<pic:cNvPicPr/>'
                f'</pic:nvPicPr>'
                f'<pic:blipFill>'
                f'<a:blip r:embed="{rid_placeholder}"/>'
                f'<a:stretch><a:fillRect/></a:stretch>'
                f'</pic:blipFill>'
                f'<pic:spPr>'
                f'<a:xfrm><a:off x="0" y="0"/><a:ext cx="{cx}" cy="{cy}"/></a:xfrm>'
                f'<a:prstGeom prst="rect"><a:avLst/></a:prstGeom>'
                f'</pic:spPr>'
                f'</pic:pic>'
                f'</a:graphicData>'
                f'</a:graphic>'
                f'</wp:inline>'
                f'</w:drawing></w:r>'
                f'</w:p>'
            )
            return etree.fromstring(xml)

        def _build_caption_paragraph(self, caption_text):
            """Italic centered caption paragraph below an image."""
            p = etree.Element(f'{{{WNS}}}p')
            pPr = etree.SubElement(p, f'{{{WNS}}}pPr')
            jc = etree.SubElement(pPr, f'{{{WNS}}}jc')
            jc.set(f'{{{WNS}}}val', 'center')
            r = etree.SubElement(p, f'{{{WNS}}}r')
            rPr = etree.SubElement(r, f'{{{WNS}}}rPr')
            rFonts = etree.SubElement(rPr, f'{{{WNS}}}rFonts')
            rFonts.set(f'{{{WNS}}}ascii', 'Trebuchet MS')
            rFonts.set(f'{{{WNS}}}hAnsi', 'Trebuchet MS')
            etree.SubElement(rPr, f'{{{WNS}}}i')  # italic
            sz = etree.SubElement(rPr, f'{{{WNS}}}sz')
            sz.set(f'{{{WNS}}}val', '18')  # 9pt
            t = etree.SubElement(r, f'{{{WNS}}}t')
            t.text = unescape(caption_text).strip()
            t.set(XML_SPACE, 'preserve')
            return p

        def _build_docx_table(self):
            """Build a w:tbl element from parsed rows."""
            tbl = etree.Element(f'{{{WNS}}}tbl')

            # Table properties: borders, width
            tblPr = etree.SubElement(tbl, f'{{{WNS}}}tblPr')
            tblStyle = etree.SubElement(tblPr, f'{{{WNS}}}tblStyle')
            tblStyle.set(f'{{{WNS}}}val', 'TableGrid')
            tblW = etree.SubElement(tblPr, f'{{{WNS}}}tblW')
            tblW.set(f'{{{WNS}}}w', '5000')
            tblW.set(f'{{{WNS}}}type', 'pct')
            # Borders
            tblBorders = etree.SubElement(tblPr, f'{{{WNS}}}tblBorders')
            for bname in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
                b = etree.SubElement(tblBorders, f'{{{WNS}}}{bname}')
                b.set(f'{{{WNS}}}val', 'single')
                b.set(f'{{{WNS}}}sz', '4')
                b.set(f'{{{WNS}}}space', '0')
                b.set(f'{{{WNS}}}color', '999999')

            for cells, is_header in self._table_rows:
                tr = etree.SubElement(tbl, f'{{{WNS}}}tr')
                for cell_text in cells:
                    tc = etree.SubElement(tr, f'{{{WNS}}}tc')
                    # Cell properties
                    tcPr = etree.SubElement(tc, f'{{{WNS}}}tcPr')
                    if is_header:
                        shd = etree.SubElement(tcPr, f'{{{WNS}}}shd')
                        shd.set(f'{{{WNS}}}val', 'clear')
                        shd.set(f'{{{WNS}}}fill', 'F0F0F0')
                    # Paragraph inside cell
                    p = etree.SubElement(tc, f'{{{WNS}}}p')
                    pPr_cell = _make_default_pPr()
                    # Remove left indent for table cells
                    ind = pPr_cell.find(f'{{{WNS}}}ind')
                    if ind is not None:
                        pPr_cell.remove(ind)
                    p.insert(0, pPr_cell)
                    r = etree.SubElement(p, f'{{{WNS}}}r')
                    rPr_cell = _make_default_rPr()
                    if is_header:
                        bold = etree.SubElement(rPr_cell, f'{{{WNS}}}b')
                    r.insert(0, rPr_cell)
                    t = etree.SubElement(r, f'{{{WNS}}}t')
                    t.text = cell_text
                    t.set(XML_SPACE, 'preserve')
            return tbl

    parser = DocxHTMLParser()
    parser.feed(html_content)
    # Flush any remaining text
    parser._flush_text()
    return parser.elements


# ── Engineering documents table ───────────────────────────────

def _fill_engineering_table(body, proposal):
    """Find the engineering documents table and replace its data rows."""
    eng_docs = list(proposal.engineering_documents.all())
    if not eng_docs:
        return

    tables = body.findall(f'.//{{{WNS}}}tbl')
    for tbl in tables:
        rows = tbl.findall(f'{{{WNS}}}tr')
        if len(rows) < 2:
            continue

        # Check header row
        header_text = ''.join(
            (t.text or '') for t in rows[0].findall(f'.//{{{WNS}}}t')
        ).lower()
        if 'document' not in header_text:
            continue

        # Copy cell formatting from a data row (row[2], skipping separator)
        cell_templates = []
        if len(rows) > 2:
            data_row = rows[2]
            for tc in data_row.findall(f'{{{WNS}}}tc'):
                tcPr = tc.find(f'{{{WNS}}}tcPr')
                rPr = None
                first_run = tc.find(f'.//{{{WNS}}}r')
                if first_run is not None:
                    rPr_el = first_run.find(f'{{{WNS}}}rPr')
                    if rPr_el is not None:
                        rPr = copy.deepcopy(rPr_el)
                cell_templates.append({
                    'tcPr': copy.deepcopy(tcPr) if tcPr is not None else None,
                    'rPr': rPr,
                })

        # Remove all rows except header
        for row in rows[1:]:
            tbl.remove(row)

        # Add new rows
        for doc in eng_docs:
            row = etree.SubElement(tbl, f'{{{WNS}}}tr')
            for ci, cell_text in enumerate([doc.doc_type, doc.doc_number, doc.doc_title]):
                tc = etree.SubElement(row, f'{{{WNS}}}tc')
                if ci < len(cell_templates) and cell_templates[ci]['tcPr'] is not None:
                    tc.insert(0, copy.deepcopy(cell_templates[ci]['tcPr']))
                p = etree.SubElement(tc, f'{{{WNS}}}p')
                r = etree.SubElement(p, f'{{{WNS}}}r')
                if ci < len(cell_templates) and cell_templates[ci]['rPr'] is not None:
                    r.insert(0, copy.deepcopy(cell_templates[ci]['rPr']))
                t = etree.SubElement(r, f'{{{WNS}}}t')
                t.text = cell_text
                t.set(XML_SPACE, 'preserve')
        break


# ── Image injection (post-process ZIP) ────────────────────────

def _resolve_image_path(src):
    """Resolve an image URL/path to a local filesystem path.

    Handles:
    - Absolute URLs (http://host/media/...)
    - Server-relative URLs (/media/...)
    - Page-relative URLs with ../ (../../../media/...)
    - Plain relative paths (media/...)
    """
    if not src:
        return None
    # Strip origin if present
    if src.startswith('http://') or src.startswith('https://'):
        try:
            from urllib.parse import urlparse
            parsed = urlparse(src)
            src = parsed.path
        except Exception:
            return None

    media_prefix = settings.MEDIA_URL.strip('/')

    # Look for the media prefix anywhere in the path (handles ../../../media/...)
    if media_prefix:
        # Normalize the path separators
        src_normalized = src.replace('\\', '/')
        marker = '/' + media_prefix + '/'
        idx = src_normalized.find(marker)
        if idx >= 0:
            rel_path = src_normalized[idx + len(marker):]
            return os.path.join(str(settings.MEDIA_ROOT), rel_path.replace('/', os.sep))
        # Starts with the media prefix
        src_clean = src_normalized.lstrip('./').lstrip('/')
        if src_clean.startswith(media_prefix + '/'):
            rel_path = src_clean[len(media_prefix) + 1:]
            return os.path.join(str(settings.MEDIA_ROOT), rel_path.replace('/', os.sep))

    return None


def _inject_images(modified_files, image_registry, all_names):
    """Add image files to word/media/, update relationships + content types,
    and replace rId placeholders in document.xml with real numeric rIds."""
    rels_name = 'word/_rels/document.xml.rels'
    if rels_name not in modified_files:
        return
    rels_xml = modified_files[rels_name]
    rels_root = etree.fromstring(rels_xml)

    # Find the highest existing numeric rId so we can allocate fresh ones
    ns_rel = 'http://schemas.openxmlformats.org/package/2006/relationships'
    existing_numeric_ids = []
    for rel in rels_root.findall(f'{{{ns_rel}}}Relationship'):
        rid = rel.get('Id') or ''
        if rid.startswith('rId') and rid[3:].isdigit():
            existing_numeric_ids.append(int(rid[3:]))
    next_rid_num = (max(existing_numeric_ids) + 1) if existing_numeric_ids else 1000

    # Content types
    ct_name = '[Content_Types].xml'
    ct_xml = modified_files.get(ct_name)
    ct_root = etree.fromstring(ct_xml) if ct_xml is not None else None
    ns_ct = 'http://schemas.openxmlformats.org/package/2006/content-types'
    existing_extensions = set()
    if ct_root is not None:
        for d in ct_root.findall(f'{{{ns_ct}}}Default'):
            existing_extensions.add((d.get('Extension') or '').lower())

    ext_to_ct = {
        'png': 'image/png',
        'jpg': 'image/jpeg',
        'jpeg': 'image/jpeg',
        'gif': 'image/gif',
        'bmp': 'image/bmp',
        'svg': 'image/svg+xml',
        'webp': 'image/webp',
    }

    # Map placeholder -> real rId so we can rewrite document.xml at the end
    rid_map = {}

    for idx, img in enumerate(image_registry):
        placeholder = f'__LEAP_IMG_{idx}__'
        path = _resolve_image_path(img['src'])
        if not path or not os.path.exists(path):
            # Missing file — leave placeholder in document.xml (will error but
            # easier to debug than mysterious broken image references)
            continue

        ext = os.path.splitext(path)[1].lower().lstrip('.') or 'png'
        if ext not in ext_to_ct:
            ext = 'png'
        real_rid = f'rId{next_rid_num}'
        next_rid_num += 1
        rid_map[placeholder] = real_rid

        media_name = f'leap-image-{idx + 1}-{next_rid_num}.{ext}'
        zip_path = f'word/media/{media_name}'

        with open(path, 'rb') as f:
            modified_files[zip_path] = f.read()

        rel = etree.SubElement(rels_root, f'{{{ns_rel}}}Relationship')
        rel.set('Id', real_rid)
        rel.set('Type', 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/image')
        rel.set('Target', f'media/{media_name}')

        if ct_root is not None and ext not in existing_extensions:
            default = etree.SubElement(ct_root, f'{{{ns_ct}}}Default')
            default.set('Extension', ext)
            default.set('ContentType', ext_to_ct[ext])
            existing_extensions.add(ext)

    modified_files[rels_name] = etree.tostring(rels_root, xml_declaration=True, encoding='UTF-8', standalone=True)
    if ct_root is not None:
        modified_files[ct_name] = etree.tostring(ct_root, xml_declaration=True, encoding='UTF-8', standalone=True)

    # Replace rId placeholders in document.xml with real ones
    if rid_map:
        doc_name = 'word/document.xml'
        doc_xml = modified_files.get(doc_name)
        if doc_xml is not None:
            if isinstance(doc_xml, bytes):
                doc_text = doc_xml.decode('utf-8')
            else:
                doc_text = doc_xml
            for placeholder, real_rid in rid_map.items():
                doc_text = doc_text.replace(placeholder, real_rid)
            modified_files[doc_name] = doc_text.encode('utf-8')


# ── Main entry point ─────────────────────────────────────────

def generate_proposal_docx(proposal):
    """Generate a DOCX by cloning the template and replacing placeholders."""
    template_path = _find_template()
    if template_path is None:
        return _generate_fallback_docx(proposal)

    with open(template_path, 'rb') as f:
        template_bytes = f.read()

    rev_date = proposal.revision_date
    date_dash = rev_date.strftime('%b-%Y').upper() if rev_date else ''
    date_slash = rev_date.strftime('%b/%Y').upper() if rev_date else ''

    # ── Substring replacements (safe for long text boxes) ─────
    textbox_replacements = {}

    # Reference + doc type (header sidebar, joined as one string)
    textbox_replacements['LNUK-IRL02125070 TECHNICAL PROP'] = (
        f'{proposal.proposal_reference} {proposal.document_type.upper()[:15]}'
    )
    textbox_replacements['LNUK-IRL02125070'] = proposal.proposal_reference

    # Dates
    textbox_replacements['OCT-2025'] = date_dash
    textbox_replacements['OCT/2025'] = date_slash

    # Client and project (footer)
    textbox_replacements['MERIDIAN CONSTRUCTION'] = proposal.client_name.upper()
    textbox_replacements['PROVISIONING OF CCTV, IIS, ACS AND FTTH SYSTEM SOLUTION'] = (
        proposal.project_description.upper() if proposal.project_description else ''
    )
    textbox_replacements['PRELIMINARY - TECHNICAL PROPOSAL'] = proposal.document_type.upper()

    # Revision
    textbox_replacements['A00'] = proposal.revision

    # Region
    textbox_replacements['UNITED KINGDOM'] = proposal.get_region_display_name().upper()

    # Company entity in the header — changes with the region (LNUK -> Global,
    # LNKSA -> Arabia, LNIRL -> Ireland).
    textbox_replacements['LEAP Networks Global Ltd.'] = proposal.get_company_name()

    # ── Exact-match replacements (initials — only when the entire
    #    text box content equals the placeholder) ──────────────
    exact_replacements = {}
    if proposal.prepared_by_initials:
        exact_replacements['AJ'] = proposal.prepared_by_initials.upper()
    if proposal.checked_by_initials:
        exact_replacements['AI'] = proposal.checked_by_initials.upper()
    if proposal.approved_by_initials:
        exact_replacements['AZ'] = proposal.approved_by_initials.upper()

    # Sort by key length descending to avoid partial matches
    textbox_replacements = dict(
        sorted(textbox_replacements.items(), key=lambda x: len(x[0]), reverse=True)
    )

    # Image registry — populated by _replace_body_sections, consumed below
    image_registry = []

    # ── Process ZIP ───────────────────────────────────────────
    output = io.BytesIO()
    modified_files = {}  # filename -> bytes (for files we add/replace)

    with zipfile.ZipFile(io.BytesIO(template_bytes), 'r') as zin:
        all_names = set(zin.namelist())
        for item in zin.infolist():
            data = zin.read(item.filename)

            if item.filename == 'word/header1.xml':
                root = etree.fromstring(data)
                _strip_highlights(root)
                _replace_textboxes_in_part(root, textbox_replacements,
                                           exact_replacements)
                data = etree.tostring(root, xml_declaration=True,
                                      encoding='UTF-8', standalone=True)

            elif item.filename == 'word/footer1.xml':
                root = etree.fromstring(data)
                _strip_highlights(root)
                _replace_textboxes_in_part(root, textbox_replacements,
                                           exact_replacements)
                data = etree.tostring(root, xml_declaration=True,
                                      encoding='UTF-8', standalone=True)

            elif item.filename == 'word/document.xml':
                root = etree.fromstring(data)
                _strip_highlights(root)
                body = root.find(f'.//{{{WNS}}}body')
                _replace_cover_page(body, proposal)
                _replace_textboxes_in_part(root, textbox_replacements,
                                           exact_replacements)
                data = etree.tostring(root, xml_declaration=True,
                                      encoding='UTF-8', standalone=True)
                data = _replace_body_sections(data, proposal, image_registry)

            modified_files[item.filename] = data

        # If we collected images, inject them and update relationships + content types
        if image_registry:
            _inject_images(modified_files, image_registry, all_names)

        # Write final ZIP
        with zipfile.ZipFile(output, 'w', zipfile.ZIP_DEFLATED) as zout:
            written = set()
            for item in zin.infolist():
                zout.writestr(item, modified_files.get(item.filename, zin.read(item.filename)))
                written.add(item.filename)
            # New files (images)
            for name, data in modified_files.items():
                if name not in written:
                    zout.writestr(name, data)

    output.seek(0)
    safe_ref = slugify(str(proposal.proposal_reference or ''))[:80] or 'proposal'
    filename = f"{safe_ref}.docx"
    response = HttpResponse(
        output.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
    )
    response['Content-Disposition'] = f'attachment; filename="{filename}"'
    return response


# ── Fallback (no template) ───────────────────────────────────

def _generate_fallback_docx(proposal):
    """Generate a simple DOCX when no template is available."""
    try:
        from docx import Document
    except ImportError:
        return HttpResponse(
            'python-docx is required for DOCX export. Install it with: pip install python-docx',
            status=500,
        )

    doc = Document()
    doc.add_heading(proposal.title, level=0)
    doc.add_paragraph(f'Reference: {proposal.proposal_reference}')
    doc.add_paragraph(f'Client: {proposal.client_name}')
    doc.add_paragraph(f'Region: {proposal.get_region_display_name()}')
    doc.add_paragraph(f'Document Type: {proposal.document_type}')
    doc.add_paragraph(f'Revision: {proposal.revision}')
    if proposal.revision_date:
        doc.add_paragraph(f'Date: {proposal.revision_date.strftime("%B %Y")}')
    doc.add_paragraph(f'Prepared by: {proposal.prepared_by_initials}')
    if proposal.checked_by_initials:
        doc.add_paragraph(f'Checked by: {proposal.checked_by_initials}')
    if proposal.approved_by_initials:
        doc.add_paragraph(f'Approved by: {proposal.approved_by_initials}')

    doc.add_page_break()

    for section in proposal.sections.all():
        doc.add_heading(section.heading, level=1)
        # Strip HTML tags for the plain fallback; the templated exporter
        # handles rich content properly.
        text = re.sub(r'<[^>]+>', '', section.content or '')
        from html import unescape
        for line in unescape(text).split('\n'):
            if line.strip():
                doc.add_paragraph(line.strip())

    eng_docs = list(proposal.engineering_documents.all())
    if eng_docs:
        doc.add_heading('Engineering Documents', level=1)
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        hdr = table.rows[0].cells
        hdr[0].text = 'Type'
        hdr[1].text = 'Number'
        hdr[2].text = 'Title'
        for ed in eng_docs:
            row = table.add_row().cells
            row[0].text = ed.doc_type
            row[1].text = ed.doc_number
            row[2].text = ed.doc_title

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    safe_ref = slugify(str(proposal.proposal_reference or ''))[:80] or 'proposal'
    filename = f"{safe_ref}.docx"
    response = HttpResponse(
        buf.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
    )
    response['Content-Disposition'] = f'attachment; filename="{filename}"'
    return response
